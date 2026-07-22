from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
import os

SAVE_FOLDER = "saved_files"

os.makedirs(
    SAVE_FOLDER,
    exist_ok=True
)


def create_docx(text, filename):
    if not filename.endswith(".docx"):
        filename += ".docx"

    path = os.path.join(
        SAVE_FOLDER,
        filename
    )

    doc = Document()
    doc.add_heading(
        "AI Generated Document",
        level=1
    )
    doc.add_paragraph(text)
    doc.save(path)
    return path


def create_pdf(text, filename):
    if not filename.endswith(".pdf"):
        filename += ".pdf"

    path = os.path.join(
        SAVE_FOLDER,
        filename
    )

    doc = SimpleDocTemplate(
        path,
        pagesize=letter,
        rightMargin=40,
        leftMargin=40,
        topMargin=40,
        bottomMargin=40
    )

    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle(
        'DocTitle',
        parent=styles['Heading1'],
        fontSize=18,
        leading=22,
        textColor=colors.HexColor('#ff2400'),
        spaceAfter=15
    )

    body_style = ParagraphStyle(
        'DocBody',
        parent=styles['Normal'],
        fontSize=11,
        leading=16,
        textColor=colors.HexColor('#111111')
    )

    story = [Paragraph("AI Generated Document", title_style)]

    lines = text.split('\n')
    for line in lines:
        if line.strip():
            safe_line = line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            story.append(Paragraph(safe_line, body_style))
            story.append(Spacer(1, 6))

    doc.build(story)
    return path